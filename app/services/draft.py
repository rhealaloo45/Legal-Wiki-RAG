"""
Draft Mode - Ephemeral Context-Aware Legal Drafting.
"""
import uuid
import time
import logging
import threading
import io
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config
from services import llm
from services import wiki

logger = logging.getLogger(__name__)

# Memory Store for draft versions
DRAFT_STORE = {}
_draft_locks = {}

def _get_draft_lock(session_id: str) -> threading.Lock:
    if session_id not in _draft_locks:
        _draft_locks[session_id] = threading.Lock()
    return _draft_locks[session_id]

STANCE_INSTRUCTIONS = {
    "tata_favorable": """
Draft from Tata's perspective as the customer/principal.

Prioritize:
- broad indemnities in Tata's favour
- strong audit and inspection rights
- expansive confidentiality protections
- strict data protection and security obligations
- broad representations and warranties from the counterparty
- strong IP ownership and licensing protections
- short cure periods for counterparty breaches
- strong termination rights for Tata
- survival of key protections after termination
- liability carve-outs benefiting Tata
- compliance obligations aligned with enterprise procurement standards
- precise drafting with minimal ambiguity

Where commercially reasonable:
- avoid unnecessary mutuality
- narrow vendor exclusions
- ensure ambiguity resolves in Tata's favour

Avoid:
- vague standards
- weak enforcement language
- broad exclusions favouring the vendor
- uncapped Tata obligations
- open-ended customer liabilities
""",

    "counterparty_favorable": """
Draft from the service provider/vendor perspective.

Prioritize:
- limited and proportionate liability exposure
- narrow indemnity obligations
- strong payment protections
- broader cure periods
- protection against indirect/consequential damages
- balanced confidentiality obligations
- retention of vendor intellectual property
- commercially reasonable compliance standards
- limitations on audit scope and frequency
- predictable termination mechanisms

Avoid:
- unlimited liability
- broad customer discretion rights
- vague compliance obligations
- one-sided risk allocation
""",

    "neutral": """
Draft in a commercially balanced and negotiation-ready manner.

Prioritize:
- mutual obligations where appropriate
- commercially reasonable standards
- balanced liability allocation
- clear drafting
- operational practicality
- enforceability
- negotiation neutrality

Avoid:
- heavily one-sided risk allocation
- unnecessary ambiguity
- extreme positions favoring either party
"""
}

# Prompt Templates
#
# Shared blocks below were added after auditing two live drafts against their
# source documents: redrafting Service Agreement 2's termination clause to
# add a 30-day notice period silently turned Tata's UNILATERAL right into a
# MUTUAL one for both parties — a real, unrequested change to who holds the
# right, flagged only by a fixed, non-specific "[DRAFTING NOTE: conflicts...]"
# sentence with nothing to say WHAT conflicted. A second indemnity-clause
# redraft kept directionality correct but silently broadened which events
# triggered indemnification beyond what the source clause covered. Same
# pattern both times: an instruction to change ONE thing drifted into
# changing adjacent things nobody asked about, and the only integrity check
# was the model grading its own homework in one vague sentence baked into the
# operative text.
#
# PRESERVE_UNLESS_ASKED and REVIEW_NOTES_TABLE below exist specifically for
# that failure shape — informed by a colleague's DRAFT_SYSTEM_PROMPT, which
# handles it two ways worth adopting: (1) soften a broad/one-sided right with
# a qualifier (materiality, reasonable opinion, applicable law) instead of
# silently making it symmetric, and (2) never let a caveat live inside the
# operative clause language — it goes in a separate, structured notes section
# instead of one ambiguous inline flag.

_PRESERVE_UNLESS_ASKED = """\
- PRESERVE EXISTING ASYMMETRY UNLESS ASKED TO CHANGE IT (CRITICAL): When redrafting or modifying a clause that already exists in the wiki context, change ONLY what the prompt actually asked for — every other term, including which party holds a right, stays exactly as the source has it. A right that is unilateral in the source (e.g. only one party may terminate for convenience, only one party indemnifies the other) MUST remain unilateral in the redraft unless the prompt explicitly asks to make it mutual. Confirmed real failure: asked to add a 30-day notice period to Tata's existing one-sided termination-for-convenience right, the redraft silently extended that right to both parties — a substantive change to who holds it, never requested and not disclosed as a change. If a one-sided right risks looking excessive, soften it with a qualifier (materiality, reasonable opinion, applicable law, legal/compliance requirement) rather than making it mutual — a qualifier narrows how the right is exercised without handing it to the other side.
- DO NOT SILENTLY WIDEN SCOPE (CRITICAL): The same discipline applies to what a clause covers, not just who it favors. Confirmed real failure: asked only to add a notice-of-claim requirement to an existing one-directional indemnity clause, the redraft also broadened WHICH events trigger indemnification beyond what the source clause listed. Only add categories, triggers, or carve-outs the prompt actually requested."""

_REVIEW_NOTES_TABLE = """\
- KEEP OPERATIVE TEXT CLEAN (CRITICAL): The clause/document text itself must read as final, ready-to-paste legal language — no citations, caveats, or commentary inside it. Every caveat, source basis, or open question goes in the separate "Legal Review Notes" section below, never inline in the drafted text.
- LEGAL REVIEW NOTES (CRITICAL): After the drafted formulation(s), add a "**Legal Review Notes**" section as a Markdown table with columns: Source basis | Changed from source (if redrafting) | Fallback / negotiation point | Confirmation needed. Be SPECIFIC — "changed from source" must name the actual delta (e.g. "source Clause 8.3 gives Tata this right alone; this draft keeps it Tata-only" or, if something WAS changed, exactly what and why), not a generic "review alignment" note. If nothing changed from source, say so explicitly rather than leaving the cell blank.
- TABLE SYNTAX (CRITICAL): The table has EXACTLY two structural rows before any data — the header row and ONE delimiter row of dashes — then every row after that is real content. Never emit a second all-dashes row; a row of "---" cells is a delimiter, not a placeholder for missing data, and repeating it renders as a blank row above your real answer. Exact shape, four columns, no more and no fewer:
  | Source basis | Changed from source (if redrafting) | Fallback / negotiation point | Confirmation needed |
  | --- | --- | --- | --- |
  | Clause 8.3, Term and Termination – SA-Tata | Preserved Tata's unilateral right; added the 30-day notice period only | Consider a mutual notice cure window if requested in negotiation | Confirm no separate wind-down payment is owed beyond Services already performed |
  That second line above (the dashes) appears ONCE, immediately after the header, and nowhere else in the table."""

CLAUSE_TEMPLATE = """
You are an expert legal drafter, drafting like senior in-house counsel: clear, concise, enforceable, negotiation-ready, and not over-engineered. Draft a specific clause based on the prompt, applying the requested stance.
Preserve numbering if applicable. Capitalize defined terms. Use precise legal drafting language and avoid vague qualifiers unless explicitly defined.
Avoid duplicative provisions — consolidate repeated concepts (e.g. insolvency, survival, accrued rights, return/destruction, audit, remedies) rather than restating them across sub-clauses. Prefer compact categories with illustrative examples over exhaustive lists.
If the prompt requests a clause type the wiki context contains NO existing provision for, say so as the first line of the Legal Review Notes ("No existing [type] clause in the retrieved context — this is market-standard language, not adapted from the agreement's own wording") — you may still draft it using standard market mechanisms, this is disclosure, not a reason to refuse.
{preserve_unless_asked}
{review_notes_table}
{wiki_instructions}

STANCE INSTRUCTIONS:
{stance_instruction}

PROMPT:
{prompt}

Begin the response with: "*AI-generated draft — requires legal review.*"
"""

FULL_DOCUMENT_TEMPLATE = """
You are an expert legal drafter, drafting like senior in-house counsel: clear, concise, enforceable, negotiation-ready, and not over-engineered. Draft a full agreement based on the prompt, applying the requested stance.
Include: agreement title, parties, recitals, definitions, commercial terms, confidentiality, IP, liability, termination, general provisions, schedules where relevant, and execution block.
Formatting should resemble enterprise commercial agreements. Use Markdown.
Avoid duplicative provisions — consolidate repeated concepts (e.g. insolvency, survival, accrued rights, return/destruction, audit, remedies) into one clause rather than restating them in several.
{preserve_unless_asked}
{review_notes_table}
{wiki_instructions}

STANCE INSTRUCTIONS:
{stance_instruction}

PROMPT:
{prompt}

Begin the response with: "*AI-generated draft — requires legal review.*"
"""

COMMUNICATION_TEMPLATE = """
You are an expert legal counsel. Draft a business/legal communication based on the prompt, applying the requested stance.
Outputs should remain concise, professional, and include structured risk summaries where appropriate.
FACTS ARE A HARD LIMIT (CRITICAL): Unlike clause drafting, this communication describes an EXISTING matter or document — every factual claim about that matter (dates, figures, party names, contract terms, status) must come from the user's prompt or the wiki context. Do not invent or infer a fact that isn't stated in either. If a fact the communication needs isn't available, list it as a "confirmation needed" item rather than stating it as if verified.
{wiki_instructions}

STANCE INSTRUCTIONS:
{stance_instruction}

PROMPT:
{prompt}

Begin the response with: "*AI-generated draft — requires legal review.*"
"""

LETTER_TEMPLATE = """
You are an expert legal counsel. Draft a formal legal letter based on the prompt, applying the requested stance.
Maintain a formal legal tone. Include factual background, identify the legal basis, specify demands clearly, and include enforceable timelines.
FACTS ARE A HARD LIMIT (CRITICAL): This letter describes an EXISTING matter — every factual claim (dates, figures, party names, contract terms, prior correspondence) must come from the user's prompt or the wiki context, never invented or inferred. If a needed fact isn't available, list it as a "confirmation needed" item rather than stating it as if verified.
{wiki_instructions}

STANCE INSTRUCTIONS:
{stance_instruction}

PROMPT:
{prompt}

Begin the response with: "*AI-generated draft — requires legal review.*"
"""

TRACKER_TEMPLATE = """
You are an expert legal compliance officer. Generate an obligations/compliance tracker based on the prompt.
Produce a Markdown table. Identify obligations, deadlines, responsible parties, gaps, and high-risk items.
The tracker output should be operationally usable by legal/compliance teams.
FACTS ARE A HARD LIMIT (CRITICAL): Every obligation, deadline, or party listed must trace to the user's prompt or the wiki context — never invent an obligation or deadline that isn't actually stated in either. A gap in the source is itself a trackable row ("not specified in source — confirm"), not something to fill with an assumed value.
{wiki_instructions}

STANCE INSTRUCTIONS:
{stance_instruction}

PROMPT:
{prompt}

Begin the response with: "*AI-generated draft — requires legal review.*"
"""

def detect_stance(prompt: str) -> str:
    prompt_lower = prompt.lower()
    tata_keywords = ["tata-friendly", "tata favourable", "tata favorable", "protect tata", "in favour of tata", "customer-friendly", "customer favorable"]
    counterparty_keywords = ["service provider friendly", "vendor-friendly", "vendor favorable", "counterparty friendly", "protect vendor", "supplier favorable"]
    neutral_keywords = ["neutral", "balanced", "mutual", "objective"]
    
    for kw in tata_keywords:
        if kw in prompt_lower:
            return "tata_favorable"
    for kw in counterparty_keywords:
        if kw in prompt_lower:
            return "counterparty_favorable"
    for kw in neutral_keywords:
        if kw in prompt_lower:
            return "neutral"
            
    return "tata_favorable"

def classify_draft(prompt: str) -> str:
    sys_prompt = "Classify this legal drafting request into exactly one of these five types:\nclause | full_document | communication | letter | tracker\n\nReturn ONLY the single category word."
    try:
        res, _ = llm.ask(sys_prompt + "\n\n" + prompt, max_tokens=10)
        category = res.strip().lower().replace(".", "").replace(",", "")
        if category in ["clause", "full_document", "communication", "letter", "tracker"]:
            return category
    except Exception as e:
        logger.error(f"Draft classification failed: {e}")
    return "clause"

def get_draft_context(session_id: str, prompt: str) -> dict:
    try:
        index = wiki._load_index(session_id)
        if not index:
            return {"has_context": False, "context": "", "source_pages": []}
            
        pages = index.get("pages", {})
        if not pages:
            return {"has_context": False, "context": "", "source_pages": []}
            
        mentioned_files = wiki._detect_mentioned_files(prompt, pages)
        file_pages = wiki._pages_from_files(pages, mentioned_files) if mentioned_files else []

        pages_for_llm = pages
        if file_pages:
            if len(pages) <= 20:
                other = [t for t in pages if t not in file_pages]
                selected_titles = file_pages + other
            else:
                # _select_relevant_pages returns (titles, usage_dict) — was
                # unpacked as a bare list here, so `for t in llm_selected`
                # iterated the 2-tuple's two ELEMENTS (the list object, then
                # the dict object) instead of page titles, and neither is
                # ever `in pages`, silently producing zero supplementary
                # pages. session_id also wasn't passed through, so this
                # always paid for the BM25+LLM fallback selection instead of
                # the free pgvector search path (see _select_relevant_pages'
                # own docstring: vector search requires session_id).
                llm_selected, _usage = wiki._select_relevant_pages(pages_for_llm, prompt, session_id)
                seen = set(file_pages)
                supplementary = [t for t in llm_selected if t not in seen]
                selected_titles = file_pages + supplementary
                # Keep subset to ~8 target pages
                selected_titles = selected_titles[:8]
        else:
            if len(pages) <= 20:
                selected_titles = list(pages.keys())
            else:
                # Same tuple bug as above: `[:15]` on a 2-tuple is a no-op
                # slice returning the tuple unchanged, so `selected_titles`
                # became (titles, usage_dict) — the loop below then found
                # nothing `in pages` and get_draft_context silently returned
                # has_context=True with EMPTY content. This is the common
                # path (>20 pages, no specific document named in the
                # prompt), so drafting was losing wiki grounding entirely
                # here in production.
                llm_selected, _usage = wiki._select_relevant_pages(pages_for_llm, prompt, session_id)
                selected_titles = llm_selected[:15]

        wiki_parts = []
        source_pages = []
        for title in selected_titles:
            if title in pages:
                page = pages[title]
                content = page.get("content", "") if isinstance(page, dict) else page
                summary = page.get("summary", "") if isinstance(page, dict) else ""
                source_doc = page.get("source", "") if isinstance(page, dict) else ""
                
                prefix = ""
                if isinstance(page, dict) and page.get("contradiction_flagged"):
                    prefix = "[WARNING: CONTRADICTION FLAGGED]\n"
                    
                page_text = f"[PAGE]\nTitle: {title}\n"
                if source_doc:
                    page_text += f"Source Document: {source_doc}\n"
                if summary:
                    page_text += f"Summary: {summary}\n"
                page_text += f"Content:\n{prefix}{content}\n"
                
                wiki_parts.append(page_text)
                source_pages.append({"title": title, "source": source_doc})
                
        wiki_content = "\n".join(wiki_parts)
        wiki_content = wiki_content[:config.MAX_DRAFT_WIKI_CONTEXT_CHARS]
        return {"has_context": True, "context": wiki_content, "source_pages": source_pages}
    except Exception as e:
        logger.error(f"Failed to get draft context: {e}")
        return {"has_context": False, "context": "", "source_pages": []}

def _run_draft_job(job_id: str, session_id: str, prompt: str, use_wiki: bool = True):
    lock = _get_draft_lock(job_id)
    try:
        with lock:
            DRAFT_STORE[job_id]["status"] = "classifying"
            
        draft_type = classify_draft(prompt)
        stance = detect_stance(prompt)
        
        with lock:
            DRAFT_STORE[job_id]["status"] = "retrieving"
            DRAFT_STORE[job_id]["metadata"] = {
                "type": draft_type,
                "stance": stance
            }
            
        if use_wiki:
            context_dict = get_draft_context(session_id, prompt)
        else:
            context_dict = {"has_context": False, "context": "", "source_pages": []}
            
        has_wiki = context_dict["has_context"]
        
        with lock:
            DRAFT_STORE[job_id]["status"] = "generating"
            DRAFT_STORE[job_id]["metadata"]["has_wiki_context"] = has_wiki
            DRAFT_STORE[job_id]["metadata"]["context_generation_mode"] = "wiki" if has_wiki else "standalone"
            DRAFT_STORE[job_id]["metadata"]["source_pages"] = context_dict["source_pages"]
            DRAFT_STORE[job_id]["metadata"]["original_context"] = context_dict["context"]
            
        template_map = {
            "clause": CLAUSE_TEMPLATE,
            "full_document": FULL_DOCUMENT_TEMPLATE,
            "communication": COMMUNICATION_TEMPLATE,
            "letter": LETTER_TEMPLATE,
            "tracker": TRACKER_TEMPLATE
        }
        
        template = template_map.get(draft_type, CLAUSE_TEMPLATE)
        stance_inst = STANCE_INSTRUCTIONS.get(stance, STANCE_INSTRUCTIONS["neutral"])
        
        wiki_instructions = ""
        if has_wiki:
            wiki_instructions = f"\nRelevant Wiki Knowledge:\n{context_dict['context']}\n\nAdditional instructions:\n- Use wiki knowledge as drafting precedent\n- Reuse terminology already found in the wiki\n- Reuse defined terms when appropriate\n- Maintain consistency with existing agreements\n- Surface conflicts through drafting notes\n"
        
        # preserve_unless_asked / review_notes_table are only referenced by
        # CLAUSE_TEMPLATE and FULL_DOCUMENT_TEMPLATE's {placeholders} — passing
        # them to str.format() for a template that doesn't use them is a no-op,
        # so it's simplest to always pass both rather than branch per draft_type.
        final_prompt = template.format(
            stance_instruction=stance_inst,
            wiki_instructions=wiki_instructions,
            prompt=prompt,
            preserve_unless_asked=_PRESERVE_UNLESS_ASKED,
            review_notes_table=_REVIEW_NOTES_TABLE,
        )

        # Uncapped before — a reasoning-model call with no max_tokens has no
        # ceiling on hidden reasoning + visible output, and a full_document/
        # tracker draft could run well past a 10k-token overall budget for
        # this call. Bounded per draft type; short types get less since
        # their whole point is brevity.
        draft_max_tokens = (
            config.MAX_TOKENS_DRAFT_LONG if draft_type in ("full_document", "tracker")
            else config.MAX_TOKENS_DRAFT_SHORT
        )
        result, _ = llm.ask(final_prompt, max_tokens=draft_max_tokens)
        
        with lock:
            DRAFT_STORE[job_id]["status"] = "complete"
            DRAFT_STORE[job_id]["current_version"] = 1
            DRAFT_STORE[job_id]["versions"] = {
                1: {
                    "text": result.strip(),
                    "prompt": prompt,
                    "timestamp": time.time()
                }
            }
            
        wiki._log_event(session_id, "DRAFT_CREATE", f"Job: {job_id} | Type: {draft_type} | Stance: {stance}")
    except Exception as e:
        logger.error(f"Draft job {job_id} failed: {e}")
        with lock:
            DRAFT_STORE[job_id]["status"] = "error"
            DRAFT_STORE[job_id]["error"] = str(e)

def _run_refine_job(job_id: str, session_id: str, instruction: str):
    lock = _get_draft_lock(job_id)
    try:
        with lock:
            DRAFT_STORE[job_id]["status"] = "refining"
            current_v = DRAFT_STORE[job_id]["current_version"]
            current_text = DRAFT_STORE[job_id]["versions"][current_v]["text"]
            meta = DRAFT_STORE[job_id].get("metadata", {})
            draft_type = meta.get("type", "clause")

        # The original wiki context used to be re-sent in full on every
        # refine ("ORIGINAL WIKI CONTEXT FOR CONSISTENCY") — but that
        # grounding is already baked into current_text from the generation
        # pass; re-including it here just doubles the input cost on every
        # refine for no benefit the existing draft text doesn't already
        # provide. Dropped entirely — refine only needs the draft + the
        # instruction.
        refine_prompt = (
            "You are an expert legal editor modifying an existing draft. "
            "Preserve existing structure, numbering, unaffected clauses, and drafting notes. "
            "Apply ONLY the requested modifications, and return the COMPLETE revised draft. "
            "Do not add explanation or commentary outside of the draft text.\n\n"
            f"EXISTING DRAFT:\n{current_text}\n\n"
            f"REFINEMENT INSTRUCTIONS:\n{instruction}"
        )

        # Same per-type cap as generation — a refine echoes the whole draft
        # back, so this bounds the call to roughly 2x the type's own budget
        # (current_text ≈ up to the cap already, plus the new completion).
        refine_max_tokens = (
            config.MAX_TOKENS_DRAFT_LONG if draft_type in ("full_document", "tracker")
            else config.MAX_TOKENS_DRAFT_SHORT
        )
        result, _ = llm.ask(refine_prompt, max_tokens=refine_max_tokens)
        
        with lock:
            new_v = current_v + 1
            DRAFT_STORE[job_id]["current_version"] = new_v
            DRAFT_STORE[job_id]["versions"][new_v] = {
                "text": result.strip(),
                "prompt": instruction,
                "timestamp": time.time()
            }
            DRAFT_STORE[job_id]["status"] = "complete"
            
        wiki._log_event(session_id, "DRAFT_REFINE", f"Job: {job_id} | New Version: {new_v}")
    except Exception as e:
        logger.error(f"Refine job {job_id} failed: {e}")
        with lock:
            DRAFT_STORE[job_id]["status"] = "error"
            DRAFT_STORE[job_id]["error"] = str(e)

def export_draft_to_docx(draft_content: str) -> bytes:
    doc = Document()
    lines = draft_content.split('\n')
    
    in_table = False
    table_rows = []
    
    def process_table():
        if not table_rows:
            return
        # Simple markdown table parser
        # Remove empty rows and separator rows
        valid_rows = []
        for r in table_rows:
            if not r.strip() or set(r.strip().replace('|', '').replace('-', '')) == set():
                continue
            cols = [c.strip() for c in r.split('|') if c.strip() or len(r.split('|')) > 1]
            if not cols:
                continue
            # Remove empty first/last if it's border
            if r.strip().startswith('|') and len(cols) > 0:
                cols = cols[1:]
            if r.strip().endswith('|') and len(cols) > 0:
                cols = cols[:-1]
            valid_rows.append(cols)
            
        if valid_rows:
            num_cols = max(len(row) for row in valid_rows)
            table = doc.add_table(rows=len(valid_rows), cols=num_cols)
            table.style = 'Table Grid'
            for r_idx, r_data in enumerate(valid_rows):
                for c_idx, c_data in enumerate(r_data):
                    if c_idx < num_cols:
                        cell = table.cell(r_idx, c_idx)
                        cell.text = c_data
        table_rows.clear()
        
    def add_formatted_run(p, text):
        # Very simple bold parsing (**text**)
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)

    for line in lines:
        stripped = line.strip()
        if stripped.startswith('|'):
            in_table = True
            table_rows.append(stripped)
            continue
        elif in_table:
            process_table()
            in_table = False
            
        if not stripped:
            doc.add_paragraph()
            continue
            
        if stripped.startswith('# '):
            p = doc.add_heading(level=1)
            add_formatted_run(p, stripped[2:])
        elif stripped.startswith('## '):
            p = doc.add_heading(level=2)
            add_formatted_run(p, stripped[3:])
        elif stripped.startswith('### '):
            p = doc.add_heading(level=3)
            add_formatted_run(p, stripped[4:])
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_run(p, stripped[2:])
        elif stripped.startswith('[DRAFTING NOTE'):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.italic = True
        else:
            p = doc.add_paragraph()
            add_formatted_run(p, stripped)
            
    if in_table:
        process_table()
        
    f = io.BytesIO()
    doc.save(f)
    return f.getvalue()
