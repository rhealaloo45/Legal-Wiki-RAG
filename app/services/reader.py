"""
Shared file reading module with OCR fallback for scanned/image-based PDFs.

Extraction strategy per PDF page:
  1. Try pypdf text extraction (fast, works for digital/text-layer PDFs).
  2. If a page yields < MIN_CHARS_PER_PAGE characters, render it as an image
     via pymupdf and OCR it with pytesseract (accurate, slower).
  3. Combine all pages into a single text string.

Graceful degradation: if pytesseract or pymupdf aren't installed,
falls back to pypdf-only extraction (original behavior) and logs a warning
so the user knows scanned pages were skipped.
"""

import contextlib
import io
import os
import re
import logging

import config

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
# If pypdf extracts fewer characters than this from a single page,
# treat it as a scanned/image page and attempt OCR.
MIN_CHARS_PER_PAGE = 50

# OCR hang protection: Tesseract can hang indefinitely on a malformed/complex
# image, which — with no timeout — freezes the whole ingest worker thread and
# is a primary cause of ingestion "getting stuck". Each Tesseract invocation is
# capped at OCR_PAGE_TIMEOUT_SECS and retried up to OCR_MAX_ATTEMPTS times; if it
# still fails, that OCR attempt is skipped and ingestion moves on to the next
# PSM mode / page / document rather than blocking forever.
OCR_PAGE_TIMEOUT_SECS = int(os.getenv("OCR_PAGE_TIMEOUT_SECS", "30"))
OCR_MAX_ATTEMPTS = int(os.getenv("OCR_MAX_ATTEMPTS", "2"))

# ---------------------------------------------------------------------------
# OCR dependency check — done once at import time
# ---------------------------------------------------------------------------
_ocr_available = False
try:
    import pytesseract
    import fitz  # pymupdf
    from PIL import Image

    _ocr_available = True
    logger.info("OCR support enabled (pytesseract + pymupdf + Pillow)")
except ImportError as e:
    logger.warning(
        "OCR dependencies not fully available (%s). Scanned PDFs will have "
        "limited text extraction. Install with: pip install pymupdf pytesseract Pillow  "
        "and ensure Tesseract is installed on your system.",
        e,
    )


def configure_tesseract(cmd_path: str | None = None):
    """Set the Tesseract executable path (useful on Windows).

    Call this once at startup if Tesseract is not on PATH, e.g.:
        configure_tesseract(r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe")
    """
    if cmd_path and _ocr_available:
        # Set on both the top-level module and the submodule to ensure
        # the path is found regardless of which attribute the library
        # version checks internally.
        pytesseract.tesseract_cmd = cmd_path
        pytesseract.pytesseract.tesseract_cmd = cmd_path
        logger.info("Tesseract path set to: %s", cmd_path)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
def read_file_with_positions(file_path: str) -> dict:
    """Read a file and return text with page-level character positions.

    Returns: {text: str, page_map: [{page_num, char_start, char_end}, ...]}
    """
    ext = os.path.splitext(file_path)[1].lower()

    with _decrypted_source(file_path) as src:
        if ext == ".pdf":
            return _read_pdf_with_positions(src)

        if ext == ".docx":
            text = _read_docx(src)
        else:
            with open(src, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = text.strip()
    # .docx has no reliable page boundary without actually laying the document
    # out (Word paginates at render time; explicit page-break markers are
    # optional and frequently absent) — treated as one page, same as .txt.
    return {
        "text": text,
        "page_map": [{"page_num": 1, "char_start": 0, "char_end": len(text)}],
    }


@contextlib.contextmanager
def _decrypted_source(file_path: str):
    """Yield a readable path for a source file, decrypting it if encrypted.

    Every reader below takes a path (pypdf, PyMuPDF, python-docx all do), so
    an encrypted upload is decrypted to a temp file for the duration of the
    read and deleted immediately after — including on exception, which is why
    this is a context manager rather than two calls a caller has to pair up
    correctly. Plaintext files pass straight through with no copy.
    """
    from services import crypto

    tmp = None
    try:
        if crypto.is_encrypted_file(file_path):
            tmp = crypto.decrypt_file_to_temp(file_path)
            yield tmp or file_path
        else:
            yield file_path
    finally:
        if tmp:
            try:
                os.remove(tmp)
            except OSError as err:
                logger.warning("Could not remove decrypted temp file %s: %s", tmp, err)


def read_file(file_path: str) -> str:
    """Read text from a .txt, .pdf, or .docx file, using OCR for scanned PDF pages.

    Returns the full document text as a single string.
    """
    ext = os.path.splitext(file_path)[1].lower()

    with _decrypted_source(file_path) as src:
        if ext == ".pdf":
            text = _read_pdf(src)
        elif ext == ".docx":
            text = _read_docx(src)
        else:
            with open(src, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()

    # Collapse excessive whitespace (shared cleanup)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = text.strip()

    # Warn if extraction yielded very little text (likely a scanned/image PDF
    # where OCR failed or wasn't available)
    fname = os.path.basename(file_path)
    if len(text) < 100:
        logger.warning(
            "EXTRACTION WARNING: '%s' yielded only %d chars. "
            "This file may be a scanned/image PDF that needs better OCR, "
            "or the file may be empty/corrupted.",
            fname, len(text),
        )
    elif len(text) < 500:
        logger.warning(
            "EXTRACTION WARNING: '%s' yielded only %d chars (low). "
            "Some pages may not have been extracted correctly.",
            fname, len(text),
        )

    return text


# ---------------------------------------------------------------------------
# PDF reading with per-page OCR fallback
# ---------------------------------------------------------------------------
def _read_pdf_with_positions(file_path: str) -> dict:
    """Read PDF and return text with per-page character offsets."""
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    page_texts: list[str] = []

    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        page_texts.append(text)

    if _ocr_available:
        needs_ocr = [i for i, t in enumerate(page_texts) if len(t) < MIN_CHARS_PER_PAGE]
        if needs_ocr:
            try:
                doc = fitz.open(file_path)
                for i in needs_ocr:
                    try:
                        ocr_text = _ocr_page(doc[i])
                        if ocr_text.strip():
                            page_texts[i] = ocr_text
                    except Exception:
                        pass
                doc.close()
            except Exception:
                pass

    full_text = "\n".join(page_texts)
    full_text = re.sub(r"\n{3,}", "\n\n", full_text)
    full_text = re.sub(r"[ \t]{2,}", " ", full_text)
    full_text = full_text.strip()

    page_map = []
    offset = 0
    for i, pt in enumerate(page_texts):
        cleaned = re.sub(r"\n{3,}", "\n\n", pt)
        cleaned = re.sub(r"[ \t]{2,}", " ", cleaned).strip()
        start = full_text.find(cleaned, max(0, offset - 50)) if cleaned else offset
        if start == -1:
            start = offset
        end = start + len(cleaned)
        page_map.append({"page_num": i + 1, "char_start": start, "char_end": end})
        offset = end

    return {"text": full_text, "page_map": page_map}


def _read_pdf(file_path: str) -> str:
    """Read PDF, falling back to OCR for pages where pypdf extracts little text."""
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    num_pages = len(reader.pages)
    fname = os.path.basename(file_path)

    # --- First pass: pypdf text extraction ---
    pypdf_texts: list[str] = []
    needs_ocr: list[int] = []

    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        pypdf_texts.append(text)
        if len(text) < MIN_CHARS_PER_PAGE:
            needs_ocr.append(i)

    # If every page had enough text, we're done
    if not needs_ocr:
        return "\n".join(pypdf_texts)

    # If OCR isn't available, warn and return what we have
    if not _ocr_available:
        logger.warning(
            "%d/%d pages in %s appear to be scanned but OCR is not available. "
            "Install pymupdf, pytesseract, and Pillow for OCR support.",
            len(needs_ocr), num_pages, fname,
        )
        return "\n".join(pypdf_texts)

    # --- Second pass: OCR only the pages that need it ---
    logger.info(
        "OCR: processing %d/%d scanned pages in %s",
        len(needs_ocr), num_pages, fname,
    )

    try:
        doc = fitz.open(file_path)
        for i in needs_ocr:
            try:
                ocr_text = _ocr_page(doc[i])
                if ocr_text.strip():
                    pypdf_texts[i] = ocr_text
                    logger.debug(
                        "  OCR page %d/%d: extracted %d chars",
                        i + 1, num_pages, len(ocr_text),
                    )
            except Exception as e:
                logger.warning("  OCR failed for page %d of %s: %s", i + 1, fname, e)
        doc.close()
    except Exception as e:
        logger.error("Failed to open %s with pymupdf for OCR: %s", fname, e)

    ocr_count = sum(
        1 for i in needs_ocr if len(pypdf_texts[i]) >= MIN_CHARS_PER_PAGE
    )
    logger.info("OCR complete for %s: %d/%d pages successfully OCR'd", fname, ocr_count, len(needs_ocr))

    _enrich_structural_pages(file_path, pypdf_texts, needs_ocr, fname)

    return "\n".join(pypdf_texts)


def _enrich_structural_pages(file_path: str, page_texts: list[str],
                             already_ocred: list[int], fname: str) -> None:
    """Stage 01 — second look at table-bearing and figure-bearing pages.

    These pages extract *successfully* by character count, so nothing else in
    the pipeline flags them: a table's rows and columns are flattened into
    loose text, and a chart on a page with a caption is never looked at. Both
    fail silently, which is what makes them worth a dedicated pass.

    Off unless STRUCTURAL_VISION_ENABLED. This is the only place in ingest
    that can fire a vision call on a page that already read fine, so it does
    not turn itself on — the cost is per page across every document, and that
    is a decision to make deliberately rather than inherit.
    """
    if not getattr(config, "STRUCTURAL_VISION_ENABLED", False):
        return
    if config.OCR_ENGINE != "azure_vision":
        # Tesseract cannot do either job: it reads glyphs, so a chart yields
        # nothing and a table yields cells stripped of the layout that gave
        # them meaning. Escalating to it would spend time to learn nothing.
        logger.info("Structural vision enabled but OCR_ENGINE is %s — skipping "
                    "(Tesseract cannot read table structure or describe figures)",
                    config.OCR_ENGINE)
        return

    from services import page_classifier

    budget = getattr(config, "STRUCTURAL_VISION_MAX_PAGES", 5)
    ocred = set(already_ocred)
    candidates: list[tuple[float, int, dict]] = []

    try:
        doc = fitz.open(file_path)
    except Exception as err:
        logger.warning("Could not reopen %s for structural pass: %s", fname, err)
        return

    try:
        for i, text in enumerate(page_texts):
            if i in ocred or i >= len(doc):
                continue  # already read as an image; a second pass adds nothing
            img_count, img_ratio = page_classifier.page_image_stats(doc[i])
            verdict = page_classifier.classify_page(text, img_count, img_ratio)
            if verdict["kind"] in (page_classifier.TABLE_BEARING,
                                   page_classifier.FIGURE_BEARING):
                rank = max(verdict["table_score"], verdict["figure_score"])
                candidates.append((rank, i, verdict))

        if not candidates:
            return

        # Strongest signals first, then capped. A document that looks like
        # tables on every page is usually a false positive on its layout, and
        # spending a vision call per page to find that out is the expensive
        # way to learn it.
        candidates.sort(reverse=True, key=lambda c: c[0])
        if len(candidates) > budget:
            logger.info("%s: %d structural page(s) detected, processing the "
                        "top %d by signal strength", fname, len(candidates), budget)
        for rank, i, verdict in candidates[:budget]:
            try:
                extra = _ocr_page_azure_vision(doc[i], verdict["vision_mode"])
            except Exception as err:
                logger.warning("Structural vision failed on page %d of %s: %s",
                               i + 1, fname, err)
                continue
            if not extra.strip():
                continue
            # Appended, never substituted. The pypdf text is the verbatim
            # source of truth that quote verification checks against; replacing
            # it with a model's reading would make every quote on that page
            # unverifiable against the original.
            page_texts[i] = (
                f"{page_texts[i]}\n\n[STRUCTURED EXTRACTION — {verdict['kind']}, "
                f"{verdict['reason']}]\n{extra}"
            )
            logger.info("Stage 01: enriched page %d of %s (%s, +%d chars)",
                        i + 1, fname, verdict["kind"], len(extra))
    finally:
        doc.close()


# ---------------------------------------------------------------------------
# DOCX reading
# ---------------------------------------------------------------------------
def _iter_docx_block_items(document):
    """Yield each paragraph and table in a .docx in true document order.

    python-docx exposes .paragraphs and .tables as two separate flat lists
    with no ordering between them — reading them separately would put every
    table (schedules, signature blocks, fee tables) after all body text
    instead of where it actually sits. Walking the body XML directly and
    wrapping each <w:p>/<w:tbl> child preserves reading order.
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _read_docx(file_path: str) -> str:
    """Extract text from a .docx, preserving paragraph/table reading order.

    Table rows are rendered as pipe-separated cells (matching how
    _read_pdf's plain extraction reads a table row left-to-right) rather
    than being dropped — legal .docx files routinely put fee schedules,
    defined-term tables, and signature blocks in actual Word tables, not
    paragraphs.
    """
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    try:
        document = Document(file_path)
    except Exception as e:
        logger.error("Failed to open %s as a .docx: %s", os.path.basename(file_path), e)
        return ""

    parts: list[str] = []
    for block in _iter_docx_block_items(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                parts.append(text)
        elif isinstance(block, Table):
            for row in block.rows:
                cells = [c.text.strip() for c in row.cells]
                line = " | ".join(cells).strip(" |")
                if line:
                    parts.append(line)

    return "\n".join(parts)


def _preprocess_for_ocr(img: "Image.Image") -> "Image.Image":
    """Apply image preprocessing to improve OCR accuracy on scanned/redacted PDFs.

    Steps:
      1. Convert to grayscale
      2. Increase contrast (optional, via histogram stretching)
      3. Apply Otsu binarization for clean black/white text
      4. Remove noise via median filter

    Falls back to the original image if any step fails (e.g. if numpy is missing).
    """
    try:
        from PIL import ImageFilter, ImageOps

        # 1. Grayscale
        gray = img.convert("L")

        # 2. Auto-contrast — stretches histogram to use full 0-255 range
        contrasted = ImageOps.autocontrast(gray, cutoff=1)

        # 3. Otsu-like binarization via simple threshold
        #    (Use Pillow's built-in point operation; no numpy needed)
        threshold = _otsu_threshold(contrasted)
        binary = contrasted.point(lambda px: 255 if px > threshold else 0, mode="1")

        # 4. Light noise removal — median filter
        cleaned = binary.convert("L").filter(ImageFilter.MedianFilter(size=3))

        return cleaned
    except Exception as e:
        logger.debug("Image preprocessing failed, using raw image: %s", e)
        return img


def _otsu_threshold(gray_img: "Image.Image") -> int:
    """Compute an Otsu-like threshold from a grayscale PIL Image.

    This avoids a numpy dependency — builds a 256-bin histogram from Pillow,
    then sweeps for the threshold that minimises intra-class variance.
    """
    hist = gray_img.histogram()  # 256 entries
    total_pixels = sum(hist)
    if total_pixels == 0:
        return 128

    sum_total = sum(i * hist[i] for i in range(256))
    sum_bg = 0.0
    weight_bg = 0
    best_thresh = 128
    best_var = 0.0

    for t in range(256):
        weight_bg += hist[t]
        if weight_bg == 0:
            continue
        weight_fg = total_pixels - weight_bg
        if weight_fg == 0:
            break

        sum_bg += t * hist[t]
        mean_bg = sum_bg / weight_bg
        mean_fg = (sum_total - sum_bg) / weight_fg

        var_between = weight_bg * weight_fg * (mean_bg - mean_fg) ** 2
        if var_between > best_var:
            best_var = var_between
            best_thresh = t

    return best_thresh


def _ocr_with_retry(img, psm: str) -> str:
    """Run one Tesseract OCR pass with a hard timeout and bounded retries.

    Tesseract with no timeout can hang forever on a bad image and freeze the
    ingest worker. Each invocation is time-boxed to OCR_PAGE_TIMEOUT_SECS
    (pytesseract raises RuntimeError on timeout) and retried up to
    OCR_MAX_ATTEMPTS times; if every attempt times out or errors, returns ""
    so the caller skips this pass and continues rather than blocking.
    """
    for attempt in range(1, OCR_MAX_ATTEMPTS + 1):
        try:
            return pytesseract.image_to_string(
                img,
                lang="eng",
                config=f"--psm {psm}",
                timeout=OCR_PAGE_TIMEOUT_SECS,
            )
        except Exception as e:
            logger.warning(
                "OCR attempt %d/%d (psm=%s) failed/timed out after %ds: %s",
                attempt, OCR_MAX_ATTEMPTS, psm, OCR_PAGE_TIMEOUT_SECS, e,
            )
    logger.warning("OCR skipped for one page (psm=%s) after %d attempts", psm, OCR_MAX_ATTEMPTS)
    return ""


_VISION_OCR_PROMPT = (
    "Transcribe every word of visible text from this scanned legal document page, "
    "exactly as it appears, in reading order. Include headings, numbered clauses, "
    "table contents (row by row), and signature-block labels. Do not summarise, "
    "translate, correct spelling/grammar, or add commentary — output the raw "
    "transcription only, no preamble."
)

# Second vision mode (§ 01 stage 01). The transcription prompt above stays as
# it is for scanned text — this one is for pages whose content is structural
# or visual, where transcribing loose words in reading order destroys exactly
# the thing worth capturing.
_VISION_TABLE_PROMPT = (
    "This legal document page contains one or more tables. Reproduce each table "
    "as GitHub-flavoured Markdown, preserving the original row and column "
    "structure exactly — every header, every cell, in its own column. Keep cell "
    "values verbatim, including units, currency symbols and footnote markers. "
    "If a cell is genuinely empty, leave it empty rather than guessing what "
    "belongs there. Where a merged cell spans columns, repeat its value across "
    "the columns it covers and note the merge beneath the table. Transcribe any "
    "text outside the table normally, after the table. Output the content only, "
    "no preamble and no commentary."
)

_VISION_DESCRIBE_PROMPT = (
    "This legal document page contains a chart, diagram, figure or other visual "
    "content. Do two things, in this order.\n"
    "1. Transcribe all visible text on the page verbatim, in reading order.\n"
    "2. Then, under a line reading exactly 'VISUAL CONTENT:', describe what the "
    "visual actually shows — its type (chart, flowchart, org chart, photograph, "
    "screenshot, signature, seal, map), what it depicts, every axis label, "
    "legend entry and data label you can read, and the relationship or trend it "
    "conveys.\n"
    "Report only what is legible. If a value or label cannot be read with "
    "confidence, say so explicitly rather than estimating it — an invented data "
    "point is far worse than an acknowledged gap in a legal record."
)

_VISION_PROMPTS = {
    "transcribe": _VISION_OCR_PROMPT,
    "table": _VISION_TABLE_PROMPT,
    "describe": _VISION_DESCRIBE_PROMPT,
}


def _ocr_page_azure_vision(page, mode: str = "transcribe") -> str:
    """Render a page and read it via the Azure OpenAI vision-capable deployment.

    Used instead of Tesseract when OCR_ENGINE=azure_vision — helpful for scans
    Tesseract garbles (skew, low DPI originals, redaction artefacts) since the
    model reads the rendered image directly rather than running local OCR.

    `mode` selects the prompt: strict transcription for scanned text, table
    reconstruction for table-bearing pages, description for chart/diagram
    pages. Tesseract has no equivalent of the latter two — it reads glyphs,
    so a chart yields nothing and a table yields its cells with the layout
    that gave them meaning discarded.
    """
    import base64
    from services import llm

    mat = fitz.Matrix(300 / 72, 300 / 72)
    pix = page.get_pixmap(matrix=mat)
    image_b64 = base64.b64encode(pix.tobytes("png")).decode("ascii")
    prompt = _VISION_PROMPTS.get(mode, _VISION_OCR_PROMPT)

    try:
        text, _usage = llm.ask_vision(image_b64, prompt, max_tokens=4096, fast=True)
        return text.strip()
    except Exception as e:
        logger.warning("Azure vision (%s) failed for a page: %s", mode, e)
        return ""


def _ocr_page(page) -> str:
    """Render a single pymupdf page to an image and OCR it.

    Uses 300 DPI for high accuracy on legal documents. Routes to the Azure
    vision deployment when OCR_ENGINE=azure_vision; otherwise runs local
    Tesseract, trying multiple PSM modes and keeping the result with the most
    extracted text. Every Tesseract call is time-boxed and retried (see
    _ocr_with_retry) so a page that Tesseract hangs on is skipped instead of
    freezing the whole ingest.
    """
    if config.OCR_ENGINE == "azure_vision":
        return _ocr_page_azure_vision(page)

    # Render at 300 DPI (default PDF is 72 DPI)
    mat = fitz.Matrix(300 / 72, 300 / 72)
    pix = page.get_pixmap(matrix=mat)

    # Convert pymupdf pixmap → PIL Image
    raw_img = Image.open(io.BytesIO(pix.tobytes("png")))

    # Preprocess for better OCR
    processed_img = _preprocess_for_ocr(raw_img)

    # Try multiple PSM modes and keep the best result
    best_text = ""
    for psm in ("6", "3", "4"):  # 6=uniform block, 3=auto, 4=single column
        text = _ocr_with_retry(processed_img, psm)
        if len(text.strip()) > len(best_text.strip()):
            best_text = text

    # If preprocessed image gave poor results, try raw image as fallback
    if len(best_text.strip()) < MIN_CHARS_PER_PAGE:
        raw_text = _ocr_with_retry(raw_img, "6")
        if len(raw_text.strip()) > len(best_text.strip()):
            best_text = raw_text

    return best_text
